from docx import Document

document = Document()

document.add_paragraph(
    "员工每天标准工作时间为 8 小时，上午 9 点上班，下午 6 点下班，中午休息 1 小时。"
)

document.add_paragraph(
    "员工如果需要远程办公，应提前向直属主管提交申请，并说明远程办公原因和预计时间。"
)

document.add_paragraph(
    "员工试用期通常为 3 个月，试用期结束前由直属主管和 HR 共同进行转正评估。"
)

document.save("employee_handbook.docx")

print("Word 文档创建成功。")